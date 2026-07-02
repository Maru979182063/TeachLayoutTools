"""覆盖 DOCX 模式判定和答案剥离行为的回归测试。"""
from docx import Document

from app.docx_processor import _is_subjective_question, _question_blocks, process_docx_source, resolve_docx_mode


def test_resolve_student_refine_for_original_paper():
    """验证resolve 学生版 refine for original paper。"""
    resolution = resolve_docx_mode(
        "2026年高考江西卷化学高考真题（原卷版）",
        {"version": "student"},
        paragraph_texts=["一、选择题", "1. 下列说法正确的是（ ）", "2. 化学实验题"],
    )

    assert resolution.resolved_mode == "student_refine"
    assert "原卷版" in resolution.trigger_hits["student_source"]


def test_resolve_teacher_to_student_tail_mode():
    """验证resolve 教师版 to 学生版 尾部 模式。"""
    resolution = resolve_docx_mode(
        "2026年高考贵州卷化学高考真题（解析版）",
        {"version": "student", "mode": "teacher_to_student_trial"},
        paragraph_texts=[
            "2026年高考贵州卷化学高考真题（解析版）",
            "一、选择题",
            "1. 题干",
            "2. 题干",
            "3. 题干",
            "4. 题干",
            "5. 题干",
            "参考答案与试题解析",
            "1.【答案】A",
        ],
    )

    assert resolution.resolved_mode == "teacher_to_student_tail"
    assert resolution.reason == "tail_answer_section_detected"


def test_resolve_inline_parse_to_hold_mode():
    """验证resolve 行内 parse to hold 模式。"""
    resolution = resolve_docx_mode(
        "精品解析：北京市清华大学附属中学高一下学期期中物理试题（解析版）",
        {"version": "student", "mode": "teacher_to_student_trial"},
        paragraph_texts=[
            "精品解析：北京市清华大学附属中学高一下学期期中物理试题（解析版）",
            "1. 题干",
            "【答案】A",
            "【解析】略",
            "本题考点：受力分析",
            "2. 题干",
            "【答案】B",
            "【解析】略",
            "3. 题干",
            "【答案】C",
        ],
    )

    assert resolution.resolved_mode == "teacher_to_student_inline"
    assert resolution.reason == "inline_parse_markers_detected"
    assert resolution.inline_answer_marker_count >= 3


def test_resolve_notes_keep_for_handout_like_material():
    """验证resolve notes keep for 讲义 like material。"""
    resolution = resolve_docx_mode(
        "高考考前必记易错要点（109大易错）2026年高考生物二轮复习讲练测",
        {"version": "student", "mode": "teacher_to_student_trial"},
        paragraph_texts=["知识梳理", "考点一", "易错提醒", "方法总结"],
    )

    assert resolution.resolved_mode == "student_refine"
    assert resolution.reason == "notes_handout_student_refine"
    assert "讲练测" in resolution.trigger_hits["notes_handout"]


def test_process_docx_source_strips_inline_answer_blocks(tmp_path):
    """验证process DOCX source strips 行内 答案 blocks。"""
    source = tmp_path / "inline.docx"
    doc = Document()
    doc.add_paragraph("一、选择题")
    doc.add_paragraph("1. 下列说法正确的是（ ）")
    doc.add_paragraph("A. 甲 B. 乙 C. 丙 D. 丁")
    doc.add_paragraph("【答案】A")
    doc.add_paragraph("【解析】这是解析")
    doc.add_paragraph("2. 第二题题干")
    doc.add_paragraph("A. 甲 B. 乙 C. 丙 D. 丁")
    doc.add_paragraph("【答案】B")
    doc.add_paragraph("【详解】这是详解")
    doc.save(source)

    result = process_docx_source(
        source,
        "TESTINLINE",
        {"version": "student", "title": "精品解析：示例试题（解析版）", "mode": "teacher_to_student_trial"},
    )
    output_doc = Document(result.output_path)
    merged_text = "\n".join(p.text for p in output_doc.paragraphs)

    assert result.report["resolved_mode"] == "teacher_to_student_inline"
    assert result.report["student_inline_answer_removed"] is True
    assert result.report["student_inline_removed_question_count"] == 2
    assert "【答案】" not in merged_text
    assert "【解析】" not in merged_text
    assert "【详解】" not in merged_text


def test_science_non_choice_section_is_subjective():
    """验证science non choice section is subjective。"""
    doc = Document()
    doc.add_paragraph("二、非选择题：本题共4小题，共58分。")
    doc.add_paragraph("15. 双氯磺酰亚胺是关键中间体。")
    doc.add_paragraph("回答下列问题：")
    doc.add_paragraph("（1）写出化学方程式")
    blocks = _question_blocks(doc)

    assert len(blocks) == 1
    block = blocks[0]
    assert block["section_subjective"] is True
    assert _is_subjective_question(block["section_subjective"], block["number"], block["text"]) is True
