"""API smoke tests for upload, job creation, review, and artifact download behavior."""

from __future__ import annotations

import io
import time

from docx import Document
from fastapi.testclient import TestClient

from app.main import app


def make_docx() -> bytes:
    """Handle make docx."""
    doc = Document()
    doc.add_paragraph("九年级数学上学期期末真题必刷常考60题 人教版")
    for index in range(1, 8):
        doc.add_paragraph(f"{index}. 已知 x + {index} = 10，求 x 的值。")
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def wait_terminal(client: TestClient, job_id: str) -> dict:
    """Handle wait terminal."""
    for _ in range(80):
        detail = client.get(f"/api/material-jobs/{job_id}", headers={"X-User-Id": "tester"}).json()
        if detail["status"] in {"HUMAN_REVIEW_REQUIRED", "FAILED", "SUCCEEDED", "CANCELLED"}:
            return detail
        time.sleep(0.1)
    raise AssertionError("job did not reach terminal state")


def test_job_requires_human_review_before_download():
    """Verify job requires human review before download."""
    client = TestClient(app)
    upload = client.post(
        "/api/files",
        headers={"X-User-Id": "tester"},
        files={
            "file": (
                "九年级上学期期末数学真题必刷常考60题（人教版）（学生版）.docx",
                make_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"file_role": "source_material"},
    )
    assert upload.status_code == 200

    job = client.post(
        "/api/material-jobs",
        headers={"X-User-Id": "tester"},
        json={
            "source_file_ids": [upload.json()["file_id"]],
            "target": {
                "material_type": "review_handout",
                "version": "student",
                "subject": "数学",
                "grade": "九年级",
                "term": "上学期期末",
                "textbook_version": "人教版",
                "title": "九年级上学期期末真题必刷常考60题",
            },
            "options": {"max_auto_fix_count": 3},
        },
    )
    assert job.status_code == 200

    terminal = wait_terminal(client, job.json()["job_id"])
    assert terminal["status"] == "HUMAN_REVIEW_REQUIRED"
    assert terminal["next_action"] == "human_approve"

    artifacts = client.get(
        f"/api/material-jobs/{job.json()['job_id']}/artifacts",
        headers={"X-User-Id": "tester"},
    ).json()["artifacts"]
    html = [item for item in artifacts if item["type"] == "html"][-1]
    pdf = [item for item in artifacts if item["type"] == "pdf"][-1]
    assert html["download_url"] is None
    assert pdf["download_url"] is None
    assert "九年级上学期期末数学真题必刷常考60题" in html["display_name"]
    assert pdf["display_name"].endswith(".pdf")

    blocked = client.get(html["download_url"] or f"/api/artifacts/{html['artifact_id']}/download", headers={"X-User-Id": "tester"})
    assert blocked.status_code == 409

    approved = client.post(
        f"/api/material-jobs/{job.json()['job_id']}/review/approve",
        headers={"X-User-Id": "tester"},
        json={},
    )
    assert approved.status_code == 200
    assert approved.json()["status"] == "SUCCEEDED"

    downloadable = client.get(f"/api/artifacts/{html['artifact_id']}/download", headers={"X-User-Id": "tester"})
    assert downloadable.status_code == 200
    downloadable_pdf = client.get(f"/api/artifacts/{pdf['artifact_id']}/download", headers={"X-User-Id": "tester"})
    assert downloadable_pdf.status_code == 200
