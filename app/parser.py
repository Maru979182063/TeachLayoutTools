"""面向 DOCX 和 PDF 源文件的基础文本抽取、输入评估和题目解析。"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def extract_text(path: Path, file_type: str) -> tuple[str, dict]:
    """从支持的源文件中提取纯文本和轻量元数据。"""
    if file_type == "docx":
        doc = Document(path)
        chunks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    chunks.append("\t".join(row_cells))
        text = "\n".join(chunks)
        return text, {"page_count": None, "table_count": len(doc.tables), "inline_shape_count": len(doc.inline_shapes)}
    if file_type == "pdf":
        reader = PdfReader(str(path))
        pages = []
        empty_pages = 0
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if not page_text.strip():
                empty_pages += 1
            pages.append(page_text)
        return "\n".join(pages), {"page_count": len(reader.pages), "empty_pages": empty_pages}
    if file_type == "zip":
        return "", {"page_count": None}
    raise ValueError("FILE_UNSUPPORTED_TYPE")


def qualify_input(file_type: str, size_bytes: int, text: str, meta: dict) -> dict:
    """评估源文件是否适合自动处理，还是应该先进入人工审核。"""
    chars = len(text.strip())
    page_count = meta.get("page_count") or 0
    empty_pages = meta.get("empty_pages") or 0
    scanned_page_ratio = empty_pages / page_count if page_count else 0.0
    question_count = len(re.findall(r"(^|\n)\s*\d+[\.、)]", text))
    confidence = 0.35
    if chars >= 500:
        confidence += 0.25
    if question_count > 0:
        confidence += 0.2
    if scanned_page_ratio == 0:
        confidence += 0.1
    confidence = min(confidence, 0.95)

    if file_type == "zip":
        decision = "processable"
        reasons = ["source_package_uploaded"]
    elif chars == 0:
        decision = "needs_human_review"
        reasons = ["empty_extractable_text"]
    elif scanned_page_ratio > 0.5:
        decision = "needs_human_review"
        reasons = ["high_scanned_page_ratio"]
    elif chars < 500 and question_count == 0:
        decision = "needs_human_review"
        reasons = ["low_text_and_no_question_structure"]
    elif confidence < 0.75:
        decision = "partially_processable"
        reasons = ["low_confidence_but_extractable"]
    else:
        decision = "processable"
        reasons = ["text_extractable", "question_structure_detected"]

    return {
        "file_type": file_type,
        "file_size_bytes": size_bytes,
        "page_count": page_count,
        "extractable_text_chars": chars,
        "extractable_text_ratio": 1.0 if chars else 0.0,
        "detected_question_count": question_count,
        "detected_section_count": 0,
        "has_scanned_pages": scanned_page_ratio > 0,
        "scanned_page_ratio": round(scanned_page_ratio, 2),
        "has_image_only_pages": scanned_page_ratio > 0,
        "has_formula_risk": bool(re.search(r"[=＋+\-×÷√∠△]", text)),
        "has_table_risk": "\t" in text,
        "unsupported_features": [],
        "confidence": round(confidence, 2),
        "decision": decision,
        "reasons": reasons,
    }


def parse_questions(text: str) -> list[dict]:
    """把抽取文本切成较粗粒度的题块，供后续规划使用。"""
    section_pattern = re.compile(r"^\s*([一二三四五六七八九十百]+[\.．、]\s*.+?（共\s*\d+\s*小题）)\s*$")
    question_pattern = re.compile(r"^\s*(\d+)[\.．、)]\s*")
    current_section = ""
    questions = []
    current_number: int | None = None
    current_lines: list[str] = []
    current_question_section = ""

    def flush_current() -> None:
        """处理flush 当前。"""
        nonlocal current_number, current_lines, current_question_section
        if current_number is None:
            return
        content = "\n".join(current_lines).strip()
        if not content:
            current_number = None
            current_lines = []
            return
        content = re.sub(r"\n{3,}", "\n\n", content)
        questions.append(
            {
                "number": current_number,
                "section": current_question_section,
                "stem": content[:1200],
                "answer": "",
                "analysis": "",
                "type": "unknown",
            }
        )
        current_number = None
        current_lines = []

    for line in text.splitlines():
        stripped = line.strip()
        section_match = section_pattern.match(stripped)
        if section_match:
            flush_current()
            current_section = section_match.group(1).strip()
            continue

        question_match = question_pattern.match(line)
        if question_match:
            flush_current()
            current_number = int(question_match.group(1))
            current_question_section = current_section
            current_lines = [question_pattern.sub("", line, count=1).strip()]
            continue

        if current_number is not None:
            current_lines.append(line)

    flush_current()
    if not questions and text.strip():
        questions.append({"number": 1, "stem": text.strip()[:800], "answer": "", "analysis": "", "type": "unknown"})
    return questions
