"""Word-first DOCX 规范化链路，负责清理来源痕迹、判定处理模式以及师生版排版。"""
from __future__ import annotations

import base64
import html
import math
import re
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from .database import DATA_DIR


PROCESSED_DIR = DATA_DIR / "processed"
PROCESSED_DIR.mkdir(exist_ok=True)

ASSET_DIR = Path(__file__).resolve().parent / "static" / "assets" / "component-library"
COMPACT_LOGO_PATH = ASSET_DIR / "header-logo-compact.png"
FALLBACK_LOGO_PATH = ASSET_DIR / "background-logo.png"
LOGO_PATH = COMPACT_LOGO_PATH if COMPACT_LOGO_PATH.exists() else FALLBACK_LOGO_PATH
GENERATED_HEADER_DIR = ASSET_DIR / "generated-headers"
GENERATED_HEADER_DIR.mkdir(parents=True, exist_ok=True)

BRAND_PATTERNS = [
    re.compile("\u5b66\u79d1\u7f51\uff08\u5317\u4eac\uff09\u80a1\u4efd\u6709\u9650\u516c\u53f8"),
    re.compile("\u5b66\u79d1\u7f51.*?\u80a1\u4efd\u6709\u9650\u516c\u53f8"),
    re.compile("\u5b66\u79d1\u7f51"),
    re.compile("www\\.zxxk\\.com", re.I),
    re.compile("zxxk\\.com", re.I),
    re.compile("\u4e0a\u597d\u6bcf\u4e00\u5802\u8bfe"),
    re.compile("\u83c1\u4f18\u7f51"),
    re.compile("\u58f0\u660e\uff1a.*", re.S),
    re.compile("\u672a\u7ecf\u4e66\u9762\u540c\u610f.*", re.S),
]

ANSWER_SECTION_MARKERS = [
    "\u53c2\u8003\u7b54\u6848\u4e0e\u8bd5\u9898\u89e3\u6790",
    "\u53c2\u8003\u7b54\u6848",
    "\u8bd5\u9898\u89e3\u6790",
    "\u7b54\u6848\u4e0e\u89e3\u6790",
    "\u3010\u7b54\u6848\u3011",
    "\u3010\u89e3\u7b54\u3011",
]

QUESTION_RE = re.compile(r"^\s*(\d{1,3})\s*[\.\uff0e\u3001]\s*")
SUBJECTIVE_SECTION_RE = re.compile(
    "\u975e\u9009\u62e9\u9898|\u89e3\u7b54\u9898|\u8ba1\u7b97\u9898|\u8bc1\u660e\u9898|\u5e94\u7528\u9898|"
    "\u7efc\u5408\u9898|\u63a2\u7a76\u9898|\u95ee\u7b54\u9898|\u9605\u8bfb\u4e0e\u601d\u8003|"
    "\u5b9e\u9a8c\u9898|\u5de5\u827a\u6d41\u7a0b\u9898|\u5b9e\u8df5\u63a2\u7a76|\u5c3a\u89c4\u4f5c\u56fe"
)
OBJECTIVE_SECTION_RE = re.compile("\u9009\u62e9\u9898|\u586b\u7a7a\u9898|\u5224\u65ad\u9898")
SECTION_HEADING_RE = re.compile(r"^\s*[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+[\u3001\uff0e\.].*(?:\u9898|\u5c0f\u9898)")
CHOICE_OPTION_RE = re.compile(r"A[\.\uff0e\u3001].*B[\.\uff0e\u3001].*C[\.\uff0e\u3001].*D[\.\uff0e\u3001]", re.S)
SUBPART_RE = re.compile(r"[\uff08(]\s*\d+\s*[\uff09)]")
SCORE_RE = re.compile(r"[\uff08(]\s*(\d+)\s*\u5206\s*[\uff09)]")
BLANK_MARKER_PREFIX = "WORD_FIRST_ANSWER_SPACE_Q"
FILL_IN_SECTION_RE = re.compile(r"\u586b\u7a7a\u9898")
CHAPTER_HEADING_RE = re.compile(r"^\s*\u7b2c\s*[0-9\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\s*[\u7ae0\u8282]")
CHAPTER_TITLE_ONLY_RE = re.compile(r"^\s*\u7b2c\s*[0-9\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\s*\u7ae0")
KNOWLEDGE_HEADING_RE = re.compile(
    "\u9519\u56e0\u5206\u6790|\u9519\u56e0|\u6613\u9519\u9677\u9631|\u907f\u9519\u653b\u7565|\u77e5\u8bc6\u94fe\u63a5|"
    "\u65b9\u6cd5\u603b\u7ed3|\u8bfe\u7a0b\u76ee\u6807|\u5b66\u4e60\u76ee\u6807|\u77e5\u8bc6\u68b3\u7406|"
    "\u4f8b\u9898\u8bb2\u89e3|\u5f3a\u5316\u8bad\u7ec3|\u80fd\u529b\u8fdb\u9636|\u8bfe\u540e\u843d\u5b9e|"
    "\u601d\u7ef4\u6e38\u620f|\u8003\u70b9\u4e32\u8bb0|\u8bb2\u7ec3\u6d4b|\u4e13\u9898\u7ec3|\u5fc5\u8bb0|\u6613\u9519"
)
STUDENT_PAGE_CAPACITY = 29.0

MODE_ALIAS = {
    "auto": "auto",
    "teacher_keep": "teacher_keep",
    "teacher_keep_trial": "teacher_keep",
    "student_refine": "student_refine",
    "student_refine_trial": "student_refine",
    "notes_keep": "notes_keep",
    "notes_keep_trial": "notes_keep",
    "chapter_handout_keep": "chapter_handout_keep",
    "teacher_to_student_auto": "teacher_to_student_auto",
    "teacher_to_student_trial": "teacher_to_student_auto",
    "teacher_to_student_tail": "teacher_to_student_tail",
    "teacher_to_student_inline": "teacher_to_student_inline",
    "teacher_to_student_mixed_hold": "teacher_to_student_mixed_hold",
}

MODE_LABELS = {
    "teacher_keep": "教师保留版",
    "student_refine": "原卷学生版整理",
    "notes_keep": "讲义资料保留版",
    "teacher_to_student_tail": "解析后置型转学生版",
    "teacher_to_student_inline": "解析混排型转学生版",
    "teacher_to_student_mixed_hold": "解析混排型暂保留",
}

TRIGGER_GROUPS = {
    "student_source": ["原卷版", "考试版", "学生版", "原卷"],
    "teacher_source": ["解析版", "教师版", "参考答案", "全解全析", "答案版"],
    "tail_answer_section": ["参考答案与试题解析", "参考答案", "答案与解析", "试题解析", "参考解析"],
    "inline_parse_markers": ["【答案】", "【解析】", "【详解】", "【点睛】", "【解答】", "精品解析"],
    "notes_handout": ["讲练测", "易错", "必记", "知识梳理", "方法总结", "学案", "导学案", "复习提纲", "错因分析", "避错攻略", "知识链接", "考点串记", "专题练"],
    "exam_paper": ["真题", "试卷", "期中", "期末", "月考", "联考", "模拟", "冲刺卷", "开学考", "收心测", "过关卷", "必刷卷"],
}

INLINE_ANSWER_MARKERS = ["【答案】", "【解析】", "【详解】", "【解答】", "【分析】", "答案：", "解析：", "详解："]
TEACHER_RED_MARKERS = INLINE_ANSWER_MARKERS + ["【点睛】", "点睛：", "解：", "故选"]


MODE_LABELS["chapter_handout_keep"] = "\u7ae0\u8282\u8bb2\u4e49\u4fdd\u7559\u7248"
TRIGGER_GROUPS["chapter_handout"] = [
    "\u7b2c1\u7ae0",
    "\u7b2c1\u8282",
    "\u9002\u7528\u76ee\u6807",
    "\u77e5\u8bc6\u68b3\u7406",
    "\u4f8b\u9898\u8fdb\u9636",
    "\u5f3a\u5316\u8bad\u7ec3",
    "\u80fd\u529b\u8fdb\u9636",
    "\u8bfe\u540e\u7ec3\u4e60",
    "\u601d\u7ef4\u63a2\u7a76",
]
SUBJECT_ALIASES = {
    "math": "math",
    "\u6570\u5b66": "math",
    "physics": "physics",
    "\u7269\u7406": "physics",
    "chemistry": "chemistry",
    "\u5316\u5b66": "chemistry",
    "biology": "biology",
    "\u751f\u7269": "biology",
    "geography": "geography",
    "\u5730\u7406": "geography",
    "history": "history",
    "\u5386\u53f2": "history",
    "politics": "politics",
    "\u9053\u5fb7\u4e0e\u6cd5\u6cbb": "politics",
    "\u653f\u6cbb": "politics",
    "chinese": "chinese",
    "\u8bed\u6587": "chinese",
    "english": "english",
    "\u82f1\u8bed": "english",
}

SUBJECT_DISPLAY_NAMES = {
    "math": "\u6570\u5b66",
    "physics": "\u7269\u7406",
    "chemistry": "\u5316\u5b66",
    "biology": "\u751f\u7269",
    "geography": "\u5730\u7406",
    "history": "\u5386\u53f2",
    "politics": "\u653f\u6cbb",
    "chinese": "\u8bed\u6587",
    "english": "\u82f1\u8bed",
}

HANDOUT_LABEL_RULES: list[tuple[re.Pattern[str], str]] = [
    (re.compile("\u8bfe\u7a0b\u76ee\u6807|\u5b66\u4e60\u76ee\u6807|\u76ee\u6807"), "\u8bfe\u7a0b\u76ee\u6807"),
    (re.compile("\u4f8b\u9898|\u5178\u4f8b|\u8bb2\u89e3|\u793a\u4f8b"), "\u4f8b\u9898\u8bb2\u89e3"),
    (re.compile("\u7ec3\u4e60\u9898|\u4e60\u9898|\u8bad\u7ec3|\u5de9\u56fa"), "\u5f3a\u5316\u8bad\u7ec3"),
    (re.compile("\u8fdb\u9636|\u63d0\u5347|\u62d3\u5c55|\u7efc\u5408"), "\u80fd\u529b\u8fdb\u9636"),
    (re.compile("\u8bfe\u540e|\u4f5c\u4e1a|\u843d\u5b9e"), "\u8bfe\u540e\u843d\u5b9e"),
    (re.compile("\u63a2\u7a76|\u601d\u7ef4|\u6e38\u620f"), "\u601d\u7ef4\u6e38\u620f"),
    (re.compile("\u77e5\u8bc6|\u68b3\u7406|\u6982\u5ff5|\u7ed3\u6784|\u7ec4\u6210|\u8d70\u8fd1"), "\u77e5\u8bc6\u68b3\u7406"),
]

HANDOUT_LABEL_TEMPLATE_ORDER = [
    "\u8bfe\u7a0b\u76ee\u6807",
    "\u77e5\u8bc6\u68b3\u7406",
    "\u4f8b\u9898\u8bb2\u89e3",
    "\u5f3a\u5316\u8bad\u7ec3",
    "\u80fd\u529b\u8fdb\u9636",
    "\u8bfe\u540e\u843d\u5b9e",
    "\u601d\u7ef4\u6e38\u620f",
]


# 供其他工作流模块复用的公开结果对象。
@dataclass
class ProcessedDocx:
    """封装处理后 DOCX 的路径、预览 HTML 和处理报告。"""
    output_path: Path
    preview_html: str
    report: dict


@dataclass
class ModeResolution:
    """描述 DOCX 在开始排版前是如何被判定处理模式的。"""
    requested_mode: str
    resolved_mode: str
    reason: str
    trigger_hits: dict[str, list[str]]
    inline_answer_marker_count: int
    tail_answer_marker_index: int | None
    paragraph_count: int


# 模式判定辅助逻辑：决定输入是保留教师版、转成学生版，
# 还是先停下来等待人工处理。
def _normalize_mode_name(mode: str | None, version: str) -> str:
    """规范化模式 名称。"""
    normalized = MODE_ALIAS.get((mode or "").strip(), "")
    if normalized:
        return normalized
    if version == "teacher":
        return "teacher_keep"
    return "auto"


def _normalize_subject(subject: str | None) -> str:
    """规范化学科。"""
    if not subject:
        return ""
    return SUBJECT_ALIASES.get(str(subject).strip(), str(subject).strip().lower())


def _count_chapter_heading_hits(paragraph_texts: list[str]) -> int:
    """统计章节 heading hits。"""
    return sum(1 for text in paragraph_texts[:160] if CHAPTER_HEADING_RE.search((text or "").strip()))


def _count_knowledge_heading_hits(paragraph_texts: list[str]) -> int:
    """统计knowledge heading hits。"""
    return sum(1 for text in paragraph_texts[:200] if KNOWLEDGE_HEADING_RE.search((text or "").strip()))


def _display_subject(subject: str | None) -> str:
    """处理显示 学科。"""
    normalized = _normalize_subject(subject)
    return SUBJECT_DISPLAY_NAMES.get(normalized, str(subject or "").strip() or "\u8d44\u6599")


def _collect_trigger_hits(evidence_text: str) -> dict[str, list[str]]:
    """收集触发词 hits。"""
    return {
        group: [token for token in tokens if token in evidence_text]
        for group, tokens in TRIGGER_GROUPS.items()
    }


def _merge_trigger_hits(*hit_maps: dict[str, list[str]]) -> dict[str, list[str]]:
    """合并触发词 hits。"""
    merged: dict[str, list[str]] = {group: [] for group in TRIGGER_GROUPS}
    for hit_map in hit_maps:
        for group, tokens in hit_map.items():
            for token in tokens:
                if token not in merged[group]:
                    merged[group].append(token)
    return merged


def _find_tail_answer_marker_index(paragraph_texts: list[str]) -> int | None:
    """查找尾部 答案 marker index。"""
    for index, text in enumerate(paragraph_texts):
        stripped = text.strip()
        if stripped and any(marker in stripped for marker in ANSWER_SECTION_MARKERS):
            return index
    return None


def _has_tail_answer_section(paragraph_count: int, marker_index: int | None) -> bool:
    """判断是否尾部 答案 section。"""
    if marker_index is None or paragraph_count <= 0:
        return False
    return marker_index >= max(6, int(paragraph_count * 0.45))


def _count_inline_answer_markers(paragraph_texts: list[str]) -> int:
    """统计行内 答案 markers。"""
    inline_markers = INLINE_ANSWER_MARKERS + ["【点睛】"]
    return sum(1 for text in paragraph_texts if any(marker in text for marker in inline_markers))


def _looks_like_inline_answer_start(text: str) -> bool:
    """判断是否like 行内 答案 start。"""
    stripped = text.strip()
    if not stripped:
        return False
    if any(marker in stripped for marker in INLINE_ANSWER_MARKERS):
        return True
    return stripped.startswith("解：") or stripped.startswith("解:") or stripped.startswith("故选")


def resolve_docx_mode(
    source_name: str,
    target: dict,
    extracted_text: str = "",
    paragraph_texts: list[str] | None = None,
) -> ModeResolution:
    """根据文件名线索、抽取文本和段落证据选择 DOCX 处理模式。"""
    version = target.get("version") or "student"
    subject = _normalize_subject(target.get("subject"))
    requested_mode = _normalize_mode_name(target.get("mode"), version)
    paragraph_texts = paragraph_texts or []
    title_text = "\n".join(part for part in [source_name, target.get("title") or ""] if part)
    evidence_parts = [
        title_text,
        extracted_text[:6000],
        "\n".join(text for text in paragraph_texts[:180] if text.strip())[:9000],
    ]
    evidence_text = "\n".join(part for part in evidence_parts if part)
    title_trigger_hits = _collect_trigger_hits(title_text)
    content_trigger_hits = _collect_trigger_hits(evidence_text)
    trigger_hits = _merge_trigger_hits(title_trigger_hits, content_trigger_hits)
    paragraph_count = len(paragraph_texts)
    chapter_heading_hits = _count_chapter_heading_hits(paragraph_texts)
    knowledge_heading_hits = _count_knowledge_heading_hits(paragraph_texts)
    tail_answer_marker_index = _find_tail_answer_marker_index(paragraph_texts)
    has_tail_answer_section = _has_tail_answer_section(paragraph_count, tail_answer_marker_index)
    inline_answer_marker_count = _count_inline_answer_markers(paragraph_texts)
    notes_like = bool(title_trigger_hits["notes_handout"])
    student_source_like = bool(trigger_hits["student_source"])
    teacher_source_like = bool(trigger_hits["teacher_source"])
    exam_like = bool(trigger_hits["exam_paper"])
    chapter_handout_like = bool(trigger_hits["chapter_handout"]) or chapter_heading_hits >= 2 or knowledge_heading_hits >= 3
    knowledge_subject = subject in {"biology", "geography", "history", "politics", "chinese", "english"}
    mixed_inline_parse = inline_answer_marker_count >= 3 and not has_tail_answer_section

    if requested_mode == "teacher_keep":
        return ModeResolution(
            requested_mode=requested_mode,
            resolved_mode="teacher_keep",
            reason="explicit_teacher_keep",
            trigger_hits=trigger_hits,
            inline_answer_marker_count=inline_answer_marker_count,
            tail_answer_marker_index=tail_answer_marker_index,
            paragraph_count=paragraph_count,
        )
    if requested_mode == "student_refine":
        return ModeResolution(
            requested_mode=requested_mode,
            resolved_mode="student_refine",
            reason="explicit_student_refine",
            trigger_hits=trigger_hits,
            inline_answer_marker_count=inline_answer_marker_count,
            tail_answer_marker_index=tail_answer_marker_index,
            paragraph_count=paragraph_count,
        )
    if requested_mode == "notes_keep":
        return ModeResolution(
            requested_mode=requested_mode,
            resolved_mode="notes_keep",
            reason="explicit_notes_keep",
            trigger_hits=trigger_hits,
            inline_answer_marker_count=inline_answer_marker_count,
            tail_answer_marker_index=tail_answer_marker_index,
            paragraph_count=paragraph_count,
        )
    if requested_mode == "chapter_handout_keep":
        return ModeResolution(
            requested_mode=requested_mode,
            resolved_mode="chapter_handout_keep",
            reason="explicit_chapter_handout_keep",
            trigger_hits=trigger_hits,
            inline_answer_marker_count=inline_answer_marker_count,
            tail_answer_marker_index=tail_answer_marker_index,
            paragraph_count=paragraph_count,
        )
    if requested_mode == "teacher_to_student_tail":
        return ModeResolution(
            requested_mode=requested_mode,
            resolved_mode="teacher_to_student_tail",
            reason="explicit_teacher_to_student_tail",
            trigger_hits=trigger_hits,
            inline_answer_marker_count=inline_answer_marker_count,
            tail_answer_marker_index=tail_answer_marker_index,
            paragraph_count=paragraph_count,
        )
    if requested_mode == "teacher_to_student_mixed_hold":
        return ModeResolution(
            requested_mode=requested_mode,
            resolved_mode="teacher_to_student_mixed_hold",
            reason="explicit_teacher_to_student_mixed_hold",
            trigger_hits=trigger_hits,
            inline_answer_marker_count=inline_answer_marker_count,
            tail_answer_marker_index=tail_answer_marker_index,
            paragraph_count=paragraph_count,
        )

    if requested_mode == "teacher_to_student_auto":
        if has_tail_answer_section:
            resolved_mode = "teacher_to_student_tail"
            reason = "tail_answer_section_detected"
        elif mixed_inline_parse:
            resolved_mode = "teacher_to_student_inline"
            reason = "inline_parse_markers_detected"
        elif chapter_handout_like and (knowledge_subject or notes_like) and not exam_like and not student_source_like and not teacher_source_like:
            resolved_mode = "chapter_handout_keep"
            reason = "chapter_handout_structure_detected"
        elif notes_like and knowledge_subject and not exam_like:
            resolved_mode = "chapter_handout_keep"
            reason = "knowledge_subject_notes_detected"
        elif notes_like:
            resolved_mode = "student_refine"
            reason = "notes_handout_student_refine"
        elif student_source_like or exam_like:
            resolved_mode = "student_refine"
            reason = "student_exam_source_detected"
        else:
            resolved_mode = "teacher_to_student_mixed_hold"
            reason = "teacher_to_student_auto_fallback_hold"
        return ModeResolution(
            requested_mode=requested_mode,
            resolved_mode=resolved_mode,
            reason=reason,
            trigger_hits=trigger_hits,
            inline_answer_marker_count=inline_answer_marker_count,
            tail_answer_marker_index=tail_answer_marker_index,
            paragraph_count=paragraph_count,
        )

    if version != "student":
        resolved_mode = "teacher_keep"
        reason = "teacher_version_default"
    elif has_tail_answer_section and teacher_source_like:
        resolved_mode = "teacher_to_student_tail"
        reason = "tail_answer_section_detected"
    elif mixed_inline_parse:
        resolved_mode = "teacher_to_student_inline"
        reason = "inline_parse_markers_detected"
    elif chapter_handout_like and (knowledge_subject or notes_like) and not exam_like and not student_source_like and not teacher_source_like:
        resolved_mode = "chapter_handout_keep"
        reason = "chapter_handout_structure_detected"
    elif notes_like and knowledge_subject and not exam_like and not student_source_like:
        resolved_mode = "chapter_handout_keep"
        reason = "knowledge_subject_notes_detected"
    elif notes_like and not student_source_like:
        resolved_mode = "student_refine"
        reason = "notes_handout_student_refine"
    elif student_source_like or exam_like:
        resolved_mode = "student_refine"
        reason = "student_exam_source_detected"
    else:
        resolved_mode = "student_refine"
        reason = "student_default_refine"

    return ModeResolution(
        requested_mode=requested_mode,
        resolved_mode=resolved_mode,
        reason=reason,
        trigger_hits=trigger_hits,
        inline_answer_marker_count=inline_answer_marker_count,
        tail_answer_marker_index=tail_answer_marker_index,
        paragraph_count=paragraph_count,
    )


# 底层文档清理辅助函数。
def _has_drawing(paragraph: Paragraph) -> bool:
    """判断是否drawing。"""
    return bool(paragraph._p.xpath(".//w:drawing | .//w:pict"))


def _delete_paragraph(paragraph: Paragraph) -> None:
    """删除段落。"""
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _clean_brand_text(text: str) -> tuple[str, bool]:
    """处理clean brand text。"""
    cleaned = text
    changed = False
    for pattern in BRAND_PATTERNS:
        next_text = pattern.sub("", cleaned)
        if next_text != cleaned:
            changed = True
            cleaned = next_text
    cleaned = re.sub(r"\s*[|｜]\s*$", "", cleaned)
    cleaned = re.sub(r"^\s*[|｜]\s*", "", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned, changed


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """处理replace 段落 text。"""
    for run in list(paragraph.runs):
        run.text = ""
    if text:
        if paragraph.runs:
            paragraph.runs[0].text = text
        else:
            paragraph.add_run(text)


def _clean_paragraph_branding(paragraph: Paragraph) -> bool:
    """处理clean 段落 品牌痕迹。"""
    original = paragraph.text
    cleaned, changed = _clean_brand_text(original)
    if not changed:
        return False

    run_level_handled = False
    for run in paragraph.runs:
        run_cleaned, run_changed = _clean_brand_text(run.text)
        if run_changed:
            run.text = run_cleaned
            run_level_handled = True

    if not run_level_handled or any(pattern.search(paragraph.text) for pattern in BRAND_PATTERNS):
        _replace_paragraph_text(paragraph, cleaned)

    if not paragraph.text.strip() and not _has_drawing(paragraph):
        _delete_paragraph(paragraph)
    return True


def _iter_container_paragraphs(container) -> list[Paragraph]:
    """处理iter 容器 段落。"""
    paragraphs: list[Paragraph] = list(container.paragraphs)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(_iter_container_paragraphs(cell))
    return paragraphs


def _clean_branding(doc: DocumentObject) -> int:
    """处理clean 品牌痕迹。"""
    cleaned = 0
    for paragraph in list(doc.paragraphs):
        if _clean_paragraph_branding(paragraph):
            cleaned += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in _iter_container_paragraphs(cell):
                    if _clean_paragraph_branding(paragraph):
                        cleaned += 1
    for section in doc.sections:
        for part in [
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]:
            for paragraph in list(_iter_container_paragraphs(part)):
                if _clean_paragraph_branding(paragraph):
                    cleaned += 1
    return cleaned


# 页眉重建辅助函数。
def _run_has_drawing(run) -> bool:
    """执行has drawing。"""
    xml = run._element.xml
    return any(token in xml for token in ("w:drawing", "w:pict", "v:shape"))


def _strip_leading_intro_graphics(doc: DocumentObject) -> dict:
    """处理strip leading intro graphics。"""
    removed_runs = 0
    affected_paragraphs: list[int] = []
    question_started = False
    for index, paragraph in enumerate(doc.paragraphs[:10], start=1):
        text = paragraph.text.strip()
        if QUESTION_RE.match(text) or SECTION_HEADING_RE.match(text):
            question_started = True
        if question_started:
            break
        removable_runs = [run for run in paragraph.runs if _run_has_drawing(run)]
        if not removable_runs:
            continue
        for run in removable_runs:
            run._element.getparent().remove(run._element)
            removed_runs += 1
        affected_paragraphs.append(index)
    return {
        "removed_runs": removed_runs,
        "affected_paragraphs": affected_paragraphs,
        "removed": removed_runs > 0,
    }


def _clear_header(header) -> None:
    """清理页眉。"""
    for child in list(header._element):
        header._element.remove(child)


def _load_header_font(size: int):
    """加载页眉 font。"""
    try:
        from PIL import ImageFont
    except Exception:
        return None
    font_candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for candidate in font_candidates:
        if candidate.exists():
            try:
                return ImageFont.truetype(str(candidate), size=size)
            except Exception:
                continue
    return None


def _set_paragraph_border_bottom(paragraph: Paragraph, color: str = "111111", size: str = "8") -> None:
    """设置段落 border bottom。"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def _set_cell_width(cell, width_twips: int) -> None:
    """设置cell width。"""
    cell.width = Inches(width_twips / 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def _set_cell_margins(cell, top: int = 0, right: int = 0, bottom: int = 0, left: int = 0) -> None:
    """设置cell margins。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("right", right), ("bottom", bottom), ("left", left)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_width(table, width_twips: int) -> None:
    """设置table width。"""
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")


def _set_table_borders_none(table) -> None:
    """设置table borders none。"""
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tbl_borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            tbl_borders.append(border)
        border.set(qn("w:val"), "nil")


def _header_strip_cache_path(subject_text: str) -> Path:
    """处理页眉 strip cache 路径。"""
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", _normalize_subject(subject_text or "material")).strip("_") or "material"
    return GENERATED_HEADER_DIR / f"header_strip_v4_{safe}.png"


def _header_brand_asset_path(subject_text: str) -> Path | None:
    """处理页眉 brand asset 路径。"""
    normalized = _normalize_subject(subject_text or "material")
    candidates = [
        ASSET_DIR / f"header-brand-{normalized}.png",
        ASSET_DIR / f"header_brand_{normalized}.png",
        GENERATED_HEADER_DIR / f"header-brand-{normalized}.png",
        GENERATED_HEADER_DIR / f"header_brand_{normalized}.png",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _trimmed_header_brand_asset_path(subject_text: str) -> Path | None:
    """处理trimmed 页眉 brand asset 路径。"""
    asset_path = _header_brand_asset_path(subject_text)
    if not asset_path or not asset_path.exists():
        return None
    trimmed_path = GENERATED_HEADER_DIR / f"header_brand_trimmed_{_normalize_subject(subject_text or 'material')}.png"
    if trimmed_path.exists():
        return trimmed_path
    try:
        from PIL import Image
    except Exception:
        return asset_path
    try:
        image = Image.open(asset_path).convert("RGBA")
        alpha = image.getchannel("A")
        bbox = alpha.getbbox()
        if not bbox:
            return asset_path
        cropped = image.crop(bbox)
        cropped.save(trimmed_path)
        return trimmed_path
    except Exception:
        return asset_path


def _static_header_strip_path(subject_text: str) -> Path | None:
    """处理static 页眉 strip 路径。"""
    normalized = _normalize_subject(subject_text or "material")
    candidates = [
        GENERATED_HEADER_DIR / f"header_strip_{normalized}.png",
        ASSET_DIR / f"header_strip_{normalized}.png",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _build_header_strip_image(subject_text: str) -> Path | None:
    """构建页眉 strip image。"""
    cache_path = _header_strip_cache_path(subject_text)
    if cache_path.exists():
        return cache_path
    try:
        from PIL import Image, ImageDraw
    except Exception:
        return None

    brand_asset = _header_brand_asset_path(subject_text)
    if brand_asset and brand_asset.exists():
        try:
            width = 1180
            height = 210
            image = Image.new("RGBA", (width, height), (255, 255, 255, 0))
            draw = ImageDraw.Draw(image)
            brand = Image.open(brand_asset).convert("RGBA")
            brand_h = 82
            brand_w = max(1, round(brand.width * (brand_h / brand.height)))
            brand = brand.resize((brand_w, brand_h), resample=Image.Resampling.LANCZOS)

            right_margin = 20
            line_y = 112
            brand_x = max(58, width - right_margin - brand_w)
            brand_y = line_y - brand_h // 2
            line_start = 58
            line_end = max(line_start + 120, brand_x - 18)
            draw.line((line_start, line_y, line_end, line_y), fill=(88, 88, 88, 255), width=3)
            image.alpha_composite(brand, (brand_x, brand_y))
            image.save(cache_path)
            return cache_path
        except Exception:
            pass

    subject_font = _load_header_font(52)
    if subject_font is None:
        return None

    width = 1800
    height = 180
    image = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)

    content_mid_y = 82
    separator_font = _load_header_font(56) or subject_font
    subject_font = _load_header_font(60) or subject_font
    subject_bbox = draw.textbbox((0, 0), subject_text, font=subject_font)
    separator_bbox = draw.textbbox((0, 0), "|", font=separator_font)
    subject_w = subject_bbox[2] - subject_bbox[0]
    subject_h = subject_bbox[3] - subject_bbox[1]
    separator_w = separator_bbox[2] - separator_bbox[0]
    separator_h = separator_bbox[3] - separator_bbox[1]

    logo = None
    logo_w = 0
    logo_h = 0
    if LOGO_PATH.exists():
        logo = Image.open(LOGO_PATH).convert("RGBA")
        logo_h = 84
        logo_w = max(1, round(logo.width * (logo_h / logo.height)))
        logo = logo.resize((logo_w, logo_h), resample=Image.Resampling.LANCZOS)

    right_margin = 28
    gap_after_logo = 18 if logo else 0
    gap_after_separator = 20
    block_width = logo_w + gap_after_logo + separator_w + gap_after_separator + subject_w
    current_x = max(58, width - right_margin - block_width)

    line_y = 112
    line_start = 58
    line_end = current_x - 16
    draw.line((line_start, line_y, line_end, line_y), fill=(88, 88, 88, 255), width=3)

    if logo:
        logo_y = content_mid_y - logo_h // 2
        image.alpha_composite(logo, (current_x, logo_y))
        current_x += logo_w + gap_after_logo

    separator_y = content_mid_y - separator_h // 2 - 2
    subject_y = content_mid_y - subject_h // 2 - 3
    draw.text((current_x, separator_y), "|", fill=(132, 132, 132, 255), font=separator_font)
    current_x += separator_w + gap_after_separator
    draw.text((current_x, subject_y), subject_text, fill=(110, 110, 110, 255), font=subject_font)

    image.save(cache_path)
    return cache_path


def _populate_header(header, title: str, subject_text: str = "") -> None:
    """填充页眉。"""
    _clear_header(header)
    subject_key = subject_text or "\u8d44\u6599"
    brand_asset = _trimmed_header_brand_asset_path(subject_key)
    if brand_asset and brand_asset.exists():
        table = header.add_table(rows=1, cols=2, width=Inches(6.45))
        table.autofit = False
        _set_table_borders_none(table)
        _set_table_width(table, 9360)

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        _set_cell_width(left_cell, 6850)
        _set_cell_width(right_cell, 2510)
        _set_cell_margins(left_cell, 0, 0, 0, 0)
        _set_cell_margins(right_cell, 0, 0, 0, 0)
        left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        left_p = left_cell.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_p.paragraph_format.space_before = Pt(0)
        left_p.paragraph_format.space_after = Pt(0)
        left_p.add_run(" ")
        _set_paragraph_border_bottom(left_p, color="666666", size="8")

        right_p = right_cell.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.paragraph_format.space_before = Pt(0)
        right_p.paragraph_format.space_after = Pt(0)
        right_p.paragraph_format.right_indent = Pt(0)
        right_p.add_run().add_picture(str(brand_asset), width=Inches(1.70))
        return
    else:
        strip_path = _static_header_strip_path(subject_key) or _build_header_strip_image(subject_key)
    if strip_path and strip_path.exists():
        paragraph = header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run().add_picture(str(strip_path), width=Inches(6.36))
    else:
        line_p = header.add_paragraph()
        line_p.paragraph_format.space_before = Pt(0)
        line_p.paragraph_format.space_after = Pt(0)
        line_p.add_run(" ")
        _set_paragraph_border_bottom(line_p, color="666666", size="6")

        brand_p = header.add_paragraph()
        brand_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        brand_p.paragraph_format.space_before = Pt(0)
        brand_p.paragraph_format.space_after = Pt(0)
        if LOGO_PATH.exists():
            brand_p.add_run().add_picture(str(LOGO_PATH), width=Inches(0.86))
        if subject_text:
            subject_run = brand_p.add_run(f" | {subject_text}")
            subject_run.font.size = Pt(12)
            subject_run.bold = True
            subject_run.font.color.rgb = RGBColor(120, 120, 120)

    spacer = header.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.add_run(" ")


def _add_header_branding(doc: DocumentObject, title: str, subject_text: str = "") -> None:
    """添加页眉 品牌痕迹。"""
    if hasattr(doc, "settings") and hasattr(doc.settings, "odd_and_even_pages_header_footer"):
        doc.settings.odd_and_even_pages_header_footer = True
    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.header_distance = Inches(0.22)
        if section.top_margin < Inches(0.7):
            section.top_margin = Inches(0.7)
        header_parts = [
            section.header,
            section.first_page_header,
            section.even_page_header,
        ]
        for part in header_parts:
            part.is_linked_to_previous = False
            _populate_header(part, title, subject_text=subject_text)


# 章节讲义标签辅助函数。
def _set_paragraph_shading(paragraph: Paragraph, fill: str) -> None:
    """设置段落 shading。"""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_paragraph_border_box(paragraph: Paragraph, color: str = "1d4dff", size: str = "16") -> None:
    """设置段落 border box。"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right"):
        border = p_bdr.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            p_bdr.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "2")
        border.set(qn("w:color"), color)


def _looks_like_content_heading(text: str) -> bool:
    """判断是否like content heading。"""
    stripped = text.strip()
    if not stripped:
        return False
    if QUESTION_RE.match(stripped):
        return False
    if re.match(r"^[A-DＡ-Ｄ][\.\uff0e\u3001]", stripped):
        return False
    if CHOICE_OPTION_RE.search(stripped):
        return False
    if len(stripped) > 24:
        return False
    if CHAPTER_TITLE_ONLY_RE.search(stripped):
        return False
    if stripped.startswith(("A", "B", "C", "D")) and len(stripped) < 12:
        return False
    return True


def _classify_handout_label(text: str) -> str | None:
    """判定讲义 label。"""
    stripped = text.strip()
    if QUESTION_RE.match(stripped):
        return None
    if re.match(r"^[A-DＡ-Ｄ][\.\uff0e\u3001]", stripped):
        return None
    for pattern, label in HANDOUT_LABEL_RULES:
        if pattern.search(stripped):
            return label
    if _looks_like_content_heading(stripped):
        return "\u77e5\u8bc6\u68b3\u7406"
    return None


def _strip_leading_cover_for_chapter_handout(doc: DocumentObject) -> dict:
    """处理strip leading cover for 章节 讲义。"""
    paragraphs = list(doc.paragraphs)
    first_chapter_index: int | None = None
    for index, paragraph in enumerate(paragraphs):
        if CHAPTER_HEADING_RE.search(paragraph.text.strip()):
            first_chapter_index = index
            break
    if first_chapter_index is None or first_chapter_index <= 0 or first_chapter_index > 24:
        return {"removed_paragraph_count": 0, "removed_cover": False}

    removed = 0
    for paragraph in list(paragraphs[:first_chapter_index]):
        _delete_paragraph(paragraph)
        removed += 1

    refreshed = list(doc.paragraphs)
    if refreshed:
        refreshed[0].paragraph_format.page_break_before = False
    return {"removed_paragraph_count": removed, "removed_cover": removed > 0}


def _remove_blank_page_break_paragraphs(doc: DocumentObject) -> int:
    """移除留白 page break 段落。"""
    removed = 0
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip():
            continue
        xml = paragraph._p.xml
        if "w:br" not in xml:
            continue
        _delete_paragraph(paragraph)
        removed += 1
    return removed


def _extract_handout_label_templates(doc: DocumentObject) -> dict[str, object]:
    """提取讲义 label templates。"""
    templates: dict[str, object] = {}
    drawing_only_paragraphs = [
        paragraph
        for paragraph in doc.paragraphs[:24]
        if _has_drawing(paragraph) and not paragraph.text.strip()
    ]
    for label, paragraph in zip(HANDOUT_LABEL_TEMPLATE_ORDER, drawing_only_paragraphs):
        templates[label] = deepcopy(paragraph._p)
    return templates


def _insert_handout_label_before(
    paragraph: Paragraph,
    label: str,
    templates: dict[str, object] | None = None,
    *,
    allow_fallback: bool = True,
) -> bool:
    """处理insert 讲义 label before。"""
    if templates and label in templates:
        paragraph._p.addprevious(deepcopy(templates[label]))
        return True
    if not allow_fallback:
        return False
    label_p = _insert_paragraph_before(paragraph)
    label_p.paragraph_format.space_before = Pt(6)
    label_p.paragraph_format.space_after = Pt(6)
    label_p.paragraph_format.keep_with_next = True
    label_p.paragraph_format.left_indent = Inches(0.08)
    label_p.paragraph_format.right_indent = Inches(4.3)
    _set_paragraph_shading(label_p, "1D4DFF")
    _set_paragraph_border_box(label_p, color="1D4DFF", size="18")
    run = label_p.add_run(f"  {label}  ")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(255, 255, 255)
    return False


def _add_functional_handout_labels(
    doc: DocumentObject,
    templates: dict[str, object] | None = None,
    *,
    allow_fallback: bool = True,
) -> list[dict]:
    """添加functional 讲义 labels。"""
    inserted: list[dict] = []
    last_label: str | None = None
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if CHAPTER_TITLE_ONLY_RE.search(text):
            last_label = None
            continue
        label = _classify_handout_label(text)
        if not label:
            continue
        previous = paragraph._p.getprevious()
        if previous is not None:
            previous_text = "".join(previous.xpath(".//*[local-name()='t']/text()")).strip()
            if previous_text == label:
                last_label = label
                continue
        if label == last_label and not re.search("\u7ec3\u4e60|\u4e60\u9898|\u8bad\u7ec3|\u63a2\u7a76|\u4f8b\u9898", text):
            continue
        used_template = _insert_handout_label_before(paragraph, label, templates=templates, allow_fallback=allow_fallback)
        if not used_template and not allow_fallback:
            continue
        inserted.append({"label": label, "anchor_text": text[:40], "used_template": used_template})
        last_label = label
    return inserted


# 教师版与学生版内容转换辅助函数。
def _strip_student_answer_section(doc: DocumentObject) -> bool:
    """处理strip 学生版 答案 section。"""
    paragraphs = list(doc.paragraphs)
    first_title = next((p.text.strip() for p in paragraphs if p.text.strip()), "")
    marker_index: int | None = None
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if any(marker in text for marker in ANSWER_SECTION_MARKERS):
            marker_index = index
            break
    if marker_index is None:
        return False

    start_index = marker_index
    while start_index > 0 and not paragraphs[start_index - 1].text.strip():
        start_index -= 1
    if start_index > 0 and first_title and paragraphs[start_index - 1].text.strip() == first_title:
        start_index -= 1
        while start_index > 0 and not paragraphs[start_index - 1].text.strip():
            start_index -= 1

    marker_element = paragraphs[start_index]._element
    body = doc._body._element
    remove_started = False
    for child in list(body):
        if child is marker_element:
            remove_started = True
        if remove_started and child.tag != qn("w:sectPr"):
            body.remove(child)
    return True


def _strip_inline_answer_blocks(doc: DocumentObject) -> dict:
    """处理strip 行内 答案 blocks。"""
    paragraphs_to_remove: list[Paragraph] = []
    affected_questions: list[int] = []

    for block in _question_blocks(doc):
        cutoff_index: int | None = None
        for index, paragraph in enumerate(block["paragraphs"][1:], start=1):
            if _looks_like_inline_answer_start(paragraph.text):
                cutoff_index = index
                break
        if cutoff_index is None:
            continue
        affected_questions.append(block["number"])
        paragraphs_to_remove.extend(block["paragraphs"][cutoff_index:])

    removed_count = 0
    seen_ids: set[int] = set()
    for paragraph in paragraphs_to_remove:
        paragraph_id = id(paragraph._p)
        if paragraph_id in seen_ids:
            continue
        seen_ids.add(paragraph_id)
        _delete_paragraph(paragraph)
        removed_count += 1

    return {
        "removed": removed_count > 0,
        "removed_paragraph_count": removed_count,
        "affected_questions": affected_questions,
        "affected_question_count": len(affected_questions),
    }


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """处理insert 段落 after。"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _insert_paragraph_before(paragraph: Paragraph) -> Paragraph:
    """处理insert 段落 before。"""
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def _color_all_runs_red(paragraph: Paragraph) -> int:
    """处理着色 all runs red。"""
    colored = 0
    for run in paragraph.runs:
        if not run.text:
            continue
        run.font.color.rgb = RGBColor(255, 0, 0)
        colored += 1
    return colored


def _color_runs_from_marker(paragraph: Paragraph, markers: list[str]) -> int:
    """处理着色 runs from marker。"""
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return 0
    hit_indexes = [full_text.find(marker) for marker in markers if marker and full_text.find(marker) >= 0]
    if not hit_indexes:
        return 0
    marker_index = min(hit_indexes)
    colored = 0
    cursor = 0
    for run in paragraph.runs:
        run_text = run.text or ""
        next_cursor = cursor + len(run_text)
        if run_text and next_cursor > marker_index:
            run.font.color.rgb = RGBColor(255, 0, 0)
            colored += 1
        cursor = next_cursor
    return colored


def _color_teacher_answer_content(doc: DocumentObject) -> dict:
    """处理着色 教师版 答案 content。"""
    paragraphs = list(doc.paragraphs)
    if not paragraphs:
        return {"colored_paragraph_count": 0, "colored_run_count": 0}

    colored_paragraph_ids: set[int] = set()
    colored_run_count = 0

    def color_paragraph(paragraph: Paragraph, *, from_marker: bool = False) -> None:
        """处理着色 段落。"""
        nonlocal colored_run_count
        para_id = id(paragraph._p)
        if para_id in colored_paragraph_ids:
            return
        colored = _color_runs_from_marker(paragraph, TEACHER_RED_MARKERS) if from_marker else _color_all_runs_red(paragraph)
        if colored:
            colored_paragraph_ids.add(para_id)
            colored_run_count += colored

    tail_marker_index = _find_tail_answer_marker_index([paragraph.text for paragraph in paragraphs])
    for block in _question_blocks(doc):
        if tail_marker_index is not None and block["start_index"] < tail_marker_index:
            continue
        inline_started = False
        for paragraph in block["paragraphs"]:
            text = paragraph.text.strip()
            if not text:
                continue
            has_marker = any(marker in text for marker in TEACHER_RED_MARKERS)
            if has_marker:
                inline_started = True
                color_paragraph(paragraph, from_marker=True)
                continue
            if inline_started:
                if QUESTION_RE.match(text) or _is_section_heading(text):
                    break
                color_paragraph(paragraph)

    # 兜底处理：对未落入题块的独立解析段落也尝试做教师版标红。
    for index, paragraph in enumerate(paragraphs):
        if tail_marker_index is not None and index < tail_marker_index:
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        if any(marker in text for marker in TEACHER_RED_MARKERS):
            color_paragraph(paragraph, from_marker=True)

    return {
        "colored_paragraph_count": len(colored_paragraph_ids),
        "colored_run_count": colored_run_count,
    }


# 学生版留白和分页辅助函数。
def _has_existing_answer_space(paragraph: Paragraph, question_number: int) -> bool:
    """判断是否existing 答案 space。"""
    node = paragraph._p
    for _ in range(12):
        node = node.getnext()
        if node is None:
            return False
        text = "".join(node.xpath(".//*[local-name()='t']/text()"))
        if QUESTION_RE.match(text.strip()):
            return False
        if f"{BLANK_MARKER_PREFIX}{question_number}" in text:
            return True
    return False


def _estimate_answer_lines(block_text: str) -> int:
    """估算答案 行。"""
    score_match = SCORE_RE.search(block_text)
    score = int(score_match.group(1)) if score_match else 0
    proof_like = bool(re.search("\u8bc1\u660e|\u8bf4\u660e\u7406\u7531|\u6c42\u8bc1|\u5c3a\u89c4\u4f5c\u56fe|\u753b\u6811\u72b6\u56fe|\u5217\u8868", block_text))
    subpart_count = len(SUBPART_RE.findall(block_text))
    science_prompt = bool(re.search("\u56de\u7b54\u4e0b\u5217\u95ee\u9898|\u5199\u51fa|\u540d\u79f0\u662f|\u7ed3\u6784\u7b80\u5f0f|\u5316\u5b66\u65b9\u7a0b\u5f0f|\u4f5c\u7528\u4e3a", block_text))
    if score >= 14 or subpart_count >= 5 or len(block_text) >= 1400:
        return 4
    if score >= 10 or proof_like or subpart_count >= 3 or (science_prompt and len(block_text) >= 500):
        return 3
    return 2


def _is_subjective_question(section_subjective: bool, number: int, block_text: str) -> bool:
    """判断是否subjective 题目。"""
    if not block_text.strip():
        return False
    if section_subjective:
        return True
    if CHOICE_OPTION_RE.search(block_text) and len(block_text) < 240:
        return False
    if re.search("\u8bc1\u660e|\u6c42\u8bc1|\u8bf4\u660e\u7406\u7531|\u5199\u51fa\u8fc7\u7a0b|\u753b\u6811\u72b6\u56fe|\u5217\u8868|\u5c3a\u89c4\u4f5c\u56fe", block_text):
        return True
    if number >= 15 and re.search("\u56de\u7b54\u4e0b\u5217\u95ee\u9898|\u5199\u51fa|\u540d\u79f0\u662f|\u7ed3\u6784\u7b80\u5f0f|\u5316\u5b66\u65b9\u7a0b\u5f0f|\u914d\u5e73\u65b9\u7a0b\u5f0f", block_text):
        return True
    if number >= 17 and len(block_text) >= 80 and SUBPART_RE.search(block_text):
        return True
    return False


def _insert_answer_lines(doc: DocumentObject, anchor: Paragraph, question_number: int, line_count: int) -> None:
    """处理insert 答案 行。"""
    previous = anchor
    for index in range(line_count):
        paragraph = _insert_paragraph_after(previous)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = index < line_count - 1
        paragraph.add_run(" ")
        if index == 0:
            marker = paragraph.add_run(f"{BLANK_MARKER_PREFIX}{question_number}")
            marker.font.hidden = True
            marker.font.size = Pt(1)
        previous = paragraph


def _question_blocks(doc: DocumentObject) -> list[dict]:
    """处理题目 blocks。"""
    paragraphs = list(doc.paragraphs)
    blocks: list[dict] = []
    section_subjective = False
    section_kind = "objective"
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if SUBJECTIVE_SECTION_RE.search(text):
            section_subjective = True
            section_kind = "subjective"
        elif OBJECTIVE_SECTION_RE.search(text):
            section_subjective = False
            section_kind = "fill_in" if FILL_IN_SECTION_RE.search(text) else "objective"

        match = QUESTION_RE.match(text)
        if match:
            blocks.append(
                {
                    "number": int(match.group(1)),
                    "start_index": index,
                    "section_subjective": section_subjective,
                    "section_kind": section_kind,
                }
            )

    for position, block in enumerate(blocks):
        start = block["start_index"]
        next_start = blocks[position + 1]["start_index"] if position + 1 < len(blocks) else len(paragraphs)
        end = next_start - 1
        for candidate_index in range(start + 1, next_start):
            candidate_text = paragraphs[candidate_index].text.strip()
            if SECTION_HEADING_RE.search(candidate_text) or SUBJECTIVE_SECTION_RE.search(candidate_text) or OBJECTIVE_SECTION_RE.search(candidate_text):
                end = candidate_index - 1
                break
        block_paragraphs = paragraphs[start : end + 1]
        block["paragraphs"] = block_paragraphs
        block["end_index"] = end
        block["text"] = "\n".join(p.text for p in block_paragraphs if p.text.strip())
    return blocks


def _is_blank_paragraph(paragraph: Paragraph) -> bool:
    """判断是否留白 段落。"""
    return not paragraph.text.strip() and not _has_drawing(paragraph)


def _is_section_heading(text: str) -> bool:
    """判断是否section heading。"""
    return bool(
        SECTION_HEADING_RE.search(text)
        or SUBJECTIVE_SECTION_RE.search(text)
        or OBJECTIVE_SECTION_RE.search(text)
    )


def _add_fill_in_spacing(doc: DocumentObject) -> int:
    """添加fill in spacing。"""
    inserted = 0
    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text or not FILL_IN_SECTION_RE.search(text) or not _is_section_heading(text):
            continue
        if index == 0:
            continue
        previous = paragraphs[index - 1]
        if _is_blank_paragraph(previous):
            continue
        spacer = _insert_paragraph_before(paragraph)
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)
        spacer.add_run(" ")
        inserted += 1
    blocks = _question_blocks(doc)
    previous_fill_in_block: dict | None = None
    for block in blocks:
        if block.get("section_kind") != "fill_in":
            previous_fill_in_block = None
            continue
        if previous_fill_in_block is None:
            previous_fill_in_block = block
            continue
        anchor = block["paragraphs"][0]
        previous_paragraph = previous_fill_in_block["paragraphs"][-1]
        if not _is_blank_paragraph(previous_paragraph):
            spacer = _insert_paragraph_before(anchor)
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            spacer.add_run(" ")
            inserted += 1
        previous_fill_in_block = block
    return inserted


def _count_drawings(paragraph: Paragraph) -> int:
    """统计drawings。"""
    return len(paragraph._p.xpath(".//w:drawing | .//w:pict"))


def _estimate_paragraph_units(paragraph: Paragraph) -> float:
    """估算段落 units。"""
    text = " ".join(paragraph.text.split())
    if not text:
        return 6.0 if _has_drawing(paragraph) else 0.35

    units = max(1.0, math.ceil(len(text) / 24))
    if _is_section_heading(text):
        units += 0.6
    if CHOICE_OPTION_RE.search(text):
        units += 1.0
    if SUBPART_RE.search(text):
        units += 0.6
    drawing_count = _count_drawings(paragraph)
    if drawing_count:
        units += drawing_count * 5.5
    return units


def _estimate_block_units(block: dict) -> float:
    """估算block units。"""
    units = sum(_estimate_paragraph_units(paragraph) for paragraph in block["paragraphs"])
    if _is_subjective_question(block["section_subjective"], block["number"], block["text"]):
        units += 0.8 + (_estimate_answer_lines(block["text"]) * 0.95)
    return units


def _block_has_visual_density(block: dict) -> bool:
    """处理block has 视觉 density。"""
    drawing_count = sum(_count_drawings(paragraph) for paragraph in block["paragraphs"])
    return drawing_count >= 2 or any(_has_drawing(paragraph) and len(paragraph.text.strip()) > 28 for paragraph in block["paragraphs"])


def _set_paragraph_pagination(
    paragraph: Paragraph,
    *,
    keep_together: bool | None = None,
    keep_with_next: bool | None = None,
    page_break_before: bool | None = None,
) -> None:
    """设置段落 pagination。"""
    if keep_together is not None:
        paragraph.paragraph_format.keep_together = keep_together
    if keep_with_next is not None:
        paragraph.paragraph_format.keep_with_next = keep_with_next
    if page_break_before is not None:
        paragraph.paragraph_format.page_break_before = page_break_before


def _normalize_page_units(units: float) -> float:
    """规范化page units。"""
    if units <= 0:
        return 0.0
    if units < STUDENT_PAGE_CAPACITY:
        return units
    remainder = units % STUDENT_PAGE_CAPACITY
    if math.isclose(remainder, 0.0, abs_tol=0.01):
        return 0.0
    return remainder


def _apply_student_pagination(doc: DocumentObject) -> dict:
    """处理apply 学生版 pagination。"""
    paragraphs = list(doc.paragraphs)
    blocks = _question_blocks(doc)
    page_units = 0.0
    previous_end = -1
    inserted_page_breaks: list[int] = []

    for block in blocks:
        interlude = paragraphs[previous_end + 1 : block["start_index"]]
        for paragraph in interlude:
            text = paragraph.text.strip()
            if text and _is_section_heading(text):
                _set_paragraph_pagination(paragraph, keep_together=True, keep_with_next=True)
            page_units = _normalize_page_units(page_units + _estimate_paragraph_units(paragraph))

        block_units = _estimate_block_units(block)
        remaining_units = STUDENT_PAGE_CAPACITY - page_units
        is_subjective = _is_subjective_question(block["section_subjective"], block["number"], block["text"])
        would_overflow = page_units + block_units > STUDENT_PAGE_CAPACITY
        needs_page_break = False
        if previous_end >= 0 and would_overflow:
            if is_subjective and block_units >= 20.0 and remaining_units < 5.5:
                needs_page_break = True
            elif is_subjective and block_units >= 14.0 and remaining_units < 3.2:
                needs_page_break = True
            elif block.get("section_kind") == "fill_in" and remaining_units < 2.2:
                needs_page_break = True
        if block is blocks[-1]:
            needs_page_break = False

        question_anchor = block["paragraphs"][0]
        compact_block = len(block["paragraphs"]) <= 3 and block_units <= 8.5 and not is_subjective
        _set_paragraph_pagination(
            question_anchor,
            page_break_before=needs_page_break,
            keep_together=compact_block,
            keep_with_next=len(block["paragraphs"]) > 1,
        )
        for index, paragraph in enumerate(block["paragraphs"]):
            if index == 0:
                continue
            if compact_block:
                _set_paragraph_pagination(
                    paragraph,
                    keep_together=True,
                    keep_with_next=index < len(block["paragraphs"]) - 1,
                )
            else:
                _set_paragraph_pagination(
                    paragraph,
                    keep_together=False,
                    keep_with_next=False,
                )

        if needs_page_break:
            inserted_page_breaks.append(block["number"])
            page_units = 0.0

        page_units = _normalize_page_units(page_units + block_units)
        previous_end = block["end_index"]

    return {
        "page_break_before_questions": inserted_page_breaks,
        "page_break_count": len(inserted_page_breaks),
    }


def _add_student_answer_spaces(doc: DocumentObject) -> list[dict]:
    """添加学生版 答案 spaces。"""
    inserted: list[dict] = []
    blocks = _question_blocks(doc)
    for index, block in enumerate(blocks):
        text = block["text"]
        number = block["number"]
        if not _is_subjective_question(block["section_subjective"], number, text):
            continue
        paragraphs = block["paragraphs"]
        if not paragraphs:
            continue
        anchor = paragraphs[-1]
        if _has_existing_answer_space(anchor, number):
            continue
        line_count = _estimate_answer_lines(text)
        if index == len(blocks) - 1:
            line_count = 0
        if line_count <= 0:
            continue
        _insert_answer_lines(doc, anchor, number, line_count)
        inserted.append({"question_number": number, "line_count": line_count})
    return inserted


def _logo_data_url() -> str:
    """处理logo data url。"""
    if not LOGO_PATH.exists():
        return ""
    data = base64.b64encode(LOGO_PATH.read_bytes()).decode("ascii")
    return f"data:image/png;base64,{data}"


def _build_preview_html(title: str, doc: DocumentObject, report: dict) -> str:
    """构建预览 HTML。"""
    logo = _logo_data_url()
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    visible = paragraphs[:220]
    blanks = report.get("student_blank_blocks") or []
    blank_summary = "\u3001".join(f"{item['question_number']}\u9898/{item['line_count']}\u884c" for item in blanks) or "\u65e0"
    mode_label = MODE_LABELS.get(report.get("resolved_mode"), report.get("resolved_mode") or "\u672a\u77e5\u6a21\u5f0f")
    mode_reason = report.get("mode_reason") or "-"
    body = "\n".join(f"<p>{html.escape(text)}</p>" for text in visible)
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>{html.escape(title)}</title>
  <style>
    body {{
      margin: 0;
      padding: 28px 42px;
      font-family: "Microsoft YaHei", "SimSun", sans-serif;
      color: #111;
      background: #f4f6f8;
      line-height: 1.75;
    }}
    .page {{
      max-width: 820px;
      margin: 0 auto;
      background: #fff;
      padding: 34px 46px 52px;
      box-shadow: 0 1px 4px rgba(16,24,40,.12);
    }}
    .masthead {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 20px;
      border-bottom: 2px solid #111;
      padding-bottom: 8px;
      margin-bottom: 22px;
    }}
    .masthead img {{ width: 118px; height: auto; display: block; }}
    h1 {{ margin: 0; font-size: 22px; text-align: right; }}
    .summary {{
      margin: 0 0 18px;
      padding: 10px 12px;
      border-left: 3px solid #111;
      background: #f8fafc;
      font-size: 13px;
    }}
    p {{ margin: 0 0 8px; }}
  </style>
</head>
<body>
  <main class="page">
    <header class="masthead">
      {'<img src="' + logo + '" alt="logo">' if logo else '<strong></strong>'}
      <h1>{html.escape(title)}</h1>
    </header>
    <div class="summary">
      DOCX \u4e3b\u4ea7\u7269\u5df2\u751f\u6210\uff1b\u6a21\u5f0f\uff1a{html.escape(mode_label)} / \u5224\u5b9a\uff1a{html.escape(mode_reason)} / \u9898\u5757\u7559\u767d\uff1a{html.escape(blank_summary)}
    </div>
    {body}
  </main>
</body>
</html>"""


def _output_path(source_path: Path, job_id: str, version: str) -> Path:
    """构建the processed DOCX 输出 路径 for one source, 任务, and 目标 版本。"""
    safe_stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", source_path.stem).strip("_") or source_path.stem
    return PROCESSED_DIR / f"{safe_stem}_{job_id}_{version}.docx"


# workflow.py 调用的 Word-first DOCX 总入口。
def process_docx_source(source_path: Path, job_id: str, target: dict, extracted_text: str = "") -> ProcessedDocx:
    """对单个源文件执行完整的 Word-first DOCX 清理与排版流程。"""
    version = target.get("version") or "student"
    title = target.get("title") or source_path.stem
    subject_text = _display_subject(target.get("subject"))
    doc = Document(source_path)
    mode_resolution = resolve_docx_mode(
        source_name=source_path.stem,
        target=target,
        extracted_text=extracted_text,
        paragraph_texts=[paragraph.text for paragraph in doc.paragraphs],
    )

    cleaned_branding = _clean_branding(doc)
    intro_graphics_report = _strip_leading_intro_graphics(doc)
    _add_header_branding(doc, title, subject_text=subject_text)

    stripped_answers = False
    inline_strip_report = {"removed": False, "removed_paragraph_count": 0, "affected_questions": [], "affected_question_count": 0}
    student_blank_blocks: list[dict] = []
    handout_cover_report = {"removed_paragraph_count": 0, "removed_cover": False}
    handout_labels: list[dict] = []
    handout_label_templates: dict[str, object] = {}
    removed_page_break_blanks = 0
    fill_in_spacing_inserted = 0
    pagination_report = {"page_break_before_questions": [], "page_break_count": 0}
    student_processing_applied = False
    teacher_red_report = {"colored_paragraph_count": 0, "colored_run_count": 0}
    if mode_resolution.resolved_mode == "teacher_to_student_tail":
        stripped_answers = _strip_student_answer_section(doc)
        fill_in_spacing_inserted = _add_fill_in_spacing(doc)
        student_blank_blocks = _add_student_answer_spaces(doc)
        pagination_report = _apply_student_pagination(doc)
        student_processing_applied = True
    elif mode_resolution.resolved_mode == "teacher_to_student_inline":
        inline_strip_report = _strip_inline_answer_blocks(doc)
        fill_in_spacing_inserted = _add_fill_in_spacing(doc)
        student_blank_blocks = _add_student_answer_spaces(doc)
        pagination_report = _apply_student_pagination(doc)
        student_processing_applied = True
    elif mode_resolution.resolved_mode == "student_refine":
        fill_in_spacing_inserted = _add_fill_in_spacing(doc)
        student_blank_blocks = _add_student_answer_spaces(doc)
        pagination_report = _apply_student_pagination(doc)
        student_processing_applied = True
    elif mode_resolution.resolved_mode == "chapter_handout_keep":
        handout_label_templates = _extract_handout_label_templates(doc)
        handout_cover_report = _strip_leading_cover_for_chapter_handout(doc)
        removed_page_break_blanks = _remove_blank_page_break_paragraphs(doc)
        handout_labels = _add_functional_handout_labels(doc, templates=handout_label_templates, allow_fallback=False)

    if version == "teacher":
        teacher_red_report = _color_teacher_answer_content(doc)

    report = {
        "processor": "word_first_docx",
        "source": str(source_path),
        "version": version,
        "requested_mode": mode_resolution.requested_mode,
        "resolved_mode": mode_resolution.resolved_mode,
        "resolved_mode_label": MODE_LABELS.get(mode_resolution.resolved_mode, mode_resolution.resolved_mode),
        "mode_reason": mode_resolution.reason,
        "mode_trigger_hits": mode_resolution.trigger_hits,
        "inline_answer_marker_count": mode_resolution.inline_answer_marker_count,
        "tail_answer_marker_index": mode_resolution.tail_answer_marker_index,
        "paragraph_count": mode_resolution.paragraph_count,
        "cleaned_branding_paragraphs": cleaned_branding,
        "leading_intro_graphics_removed": intro_graphics_report["removed"],
        "leading_intro_graphics_removed_runs": intro_graphics_report["removed_runs"],
        "leading_intro_graphics_affected_paragraphs": intro_graphics_report["affected_paragraphs"],
        "header_rebuilt": True,
        "header_subject_text": subject_text,
        "student_answer_section_removed": stripped_answers,
        "student_inline_answer_removed": inline_strip_report["removed"],
        "student_inline_removed_paragraph_count": inline_strip_report["removed_paragraph_count"],
        "student_inline_removed_question_count": inline_strip_report["affected_question_count"],
        "student_inline_removed_questions": inline_strip_report["affected_questions"],
        "handout_cover_removed": handout_cover_report["removed_cover"],
        "handout_cover_removed_paragraph_count": handout_cover_report["removed_paragraph_count"],
        "handout_removed_page_break_paragraph_count": removed_page_break_blanks,
        "handout_label_template_count": len(handout_label_templates),
        "handout_labels": handout_labels,
        "handout_label_count": len(handout_labels),
        "student_fill_in_spacing_inserted": fill_in_spacing_inserted,
        "student_blank_blocks": student_blank_blocks,
        "student_blank_block_count": len(student_blank_blocks),
        "student_blank_line_count": sum(item["line_count"] for item in student_blank_blocks),
        "student_page_break_before_questions": pagination_report["page_break_before_questions"],
        "student_page_break_count": pagination_report["page_break_count"],
        "student_processing_applied": student_processing_applied,
        "student_conversion_blocked": mode_resolution.resolved_mode in {"notes_keep", "chapter_handout_keep", "teacher_to_student_mixed_hold"},
        "teacher_content_preserved": mode_resolution.resolved_mode in {"teacher_keep", "notes_keep", "chapter_handout_keep", "teacher_to_student_mixed_hold"},
        "teacher_red_paragraph_count": teacher_red_report["colored_paragraph_count"],
        "teacher_red_run_count": teacher_red_report["colored_run_count"],
    }

    output_path = _output_path(source_path, job_id, version)
    doc.save(output_path)
    preview_html = _build_preview_html(title, doc, report)
    return ProcessedDocx(output_path=output_path, preview_html=preview_html, report=report)
